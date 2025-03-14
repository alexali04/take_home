"""
Nice reference: https://docs.anthropic.com/en/docs/about-claude/use-case-guides/legal-summarization

https://github.com/anthropics/anthropic-cookbook/blob/4a7be656bd1a6e2149f9f1c40678dac379a857a7/skills/summarization/guide.ipynb

Options:

summarization
---- prompt model to summarize the SOP
---- disadvantage: lossy summary may lose key legal details

few-shot summarization
---- we prompt the model w/ examples of documents + regulatory clauses
---- helps the model "understand" the task
---- disadvantage: longer context lengths, may infer "bad patterns"

guided summarization (chosen)
---- provide the model w/ detailed instructions

practically, domain guided summarization
---- if we know we're dealing with a warehouse compliance doc, we'd prompt it w/ specific terminology
---- not chosen due to time constraint
"""


import anthropic
import docx
import dotenv
import json
import os

from utils.prompting import Regulatory_API_Prompt

class SOP_Processor():
    """
    Processes and chunks SOP. 
    """
    def __init__(
        self, 
        sop_path: str = "./regulation_task/data/sop/original.docx",
        cut_off: bool = False, 
        cut_off_length: int = 100
    ):
        assert sop_path.endswith(".docx"), "SOP must be a .docx file"

        self.sop_path = sop_path
        self.cut_off = cut_off
        self.cut_off_length = cut_off_length
    
    def manual_para_chunking(self):
        f = open(self.sop_path, "rb")
        doc = docx.Document(f)

        paras = []
        curr_para = ""
        in_para = False

        for para in doc.paragraphs: 
            if self.cut_off and len(curr_para) > self.cut_off_length:
                curr_paras = [curr_para[:self.cut_off_length], curr_para[self.cut_off_length:]]
                paras.extend(curr_paras)
                curr_para = ""

            if para.text.strip() == "": 
                if not in_para:
                    continue 
                else: # we are in a paragraph and we hit an empty line --> we are done with the paragraph
                    in_para = False
                    paras.append(curr_para)
                    curr_para = ""

            else: # not empty line --> we are in a paragraph
                if not in_para: 
                    in_para = True
                curr_para += para.text + "\n"
            
        return paras

    def get_text_from_docx(self):
        f = open(self.sop_path, "rb")
        doc = docx.Document(f)
        return "\n".join([para.text for para in doc.paragraphs])

    def post_process_clauses(self, clause_str: str):
        """Returns json object from string"""
        clauses = json.loads(clause_str)
        return clauses

    def extract_clauses_from_docx(self, api_prompt: Regulatory_API_Prompt, context: str = ""):
        """
        Args:
            (Regulatory_API_Prompt) api_prompt: API prompt to use for clause extraction
            (str) context: context to use for clause extraction
        
        Returns:
            (str) path to the json file containing the clauses
        """
        if context == "": 
            context = self.get_text_from_docx()
        
        api_prompt.content += context

        api_key = dotenv.get_key(".env", "ANTHROPIC_API_KEY")

        messages = [
            {"role": "user", "content": api_prompt.content},
        ]

        if api_prompt.assistant_prompt is not None:
            messages.append({"role": "assistant", "content": api_prompt.assistant_prompt})

        if api_prompt.max_tokens is None:
            api_prompt.max_tokens = (len(api_prompt.content) // 4) * 3
        

        client = anthropic.Anthropic(api_key=api_key)
        response = client.messages.create(
            model=api_prompt.model,
            max_tokens=api_prompt.max_tokens,
            system=api_prompt.sys_prompt,
            messages=messages
        )

        if context == "": 
            return self.post_process_clauses(response.content[0].text)
        else: 
            return response.content[0].text

def process_sop(args, data_dir: str):
    data_dir = args.data_dir

    sop_path = os.path.join(data_dir, "sop", args.sop_name)
    prompt_dir = os.path.join(data_dir, "prompts")

    sop_processor = SOP_Processor(
        sop_path=sop_path,
        cut_off=args.cut_off,
        cut_off_length=args.cut_off_length
    )
    
    sys_prompt = open(f"{prompt_dir}/{args.sys_prompt_text}.txt").read()
    extraction_prompt = open(f"{prompt_dir}/{args.extraction_prompt_text}.txt").read()

    api_prompt = Regulatory_API_Prompt(
        model=args.model,
        sys_prompt=sys_prompt,
        assistant_prompt=None,
        content=extraction_prompt,
        max_tokens=2000
    )

    clause_dict = sop_processor.extract_clauses_from_docx(api_prompt)
    clauses_path = os.path.join(data_dir, "sop", f"{args.sop_name}_clauses.json")
    json.dump(clause_dict, open(clauses_path, "w"))

    return clauses_path

    
    







    
