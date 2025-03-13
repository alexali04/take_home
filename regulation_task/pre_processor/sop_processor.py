"""
Nice reference: https://docs.anthropic.com/en/docs/about-claude/use-case-guides/legal-summarization

https://github.com/anthropics/anthropic-cookbook/blob/4a7be656bd1a6e2149f9f1c40678dac379a857a7/skills/summarization/guide.ipynb

Options:

summarization
---- prompt model to summarize the SOP
---- disadvantage: lossy summary may lose key legal details

few-shot summarization
---- we prompt the model w/ examples of documents + regulatory clauses
---- helps the model "understand" the task
---- disadvantage: longer context lengths, may infer "bad patterns"

guided summarization (chosen)
---- provide the model w/ detailed instructions

practically, domain guided summarization
---- if we know we're dealing with a warehouse compliance doc, we'd prompt it w/ specific terminology
---- not chosen due to time constraint
"""


import anthropic
from dataclasses import dataclass
import docx
import dotenv
from typing import Optional

@dataclass
class Regulatory_API_Prompt():
    """
    Defines API call for regulatory clause extraction.  
    """
    model: str = "claude-3-5-haiku-20241022"
    sys_prompt: str = "You are a regulatory compliance expert. You are given a document and your job is to extract a numbered list of regulatory clauses from the document."
    user_prompt: str = "Document: "
    max_tokens: int = 100
    assistant_prompt: Optional[str] = None



class SOP_Processor():
    """
    Processes and chunks SOP. 
    """
    def __init__(
        self, 
        sop_path: str = "./regulation_task/data/sop/original.docx",
        cut_off: bool = False, 
        cut_off_length: int = 100
    ):
        assert sop_path.endswith(".docx"), "SOP must be a .docx file"

        self.sop_path = sop_path
        self.cut_off = cut_off
        self.cut_off_length = cut_off_length
    
    def manual_para_chunking(self):
        f = open(self.sop_path, "rb")
        doc = docx.Document(f)

        paras = []
        curr_para = ""
        in_para = False

        for para in doc.paragraphs: 
            if self.cut_off and len(curr_para) > self.cut_off_length:
                curr_paras = [curr_para[:self.cut_off_length], curr_para[self.cut_off_length:]]
                paras.extend(curr_paras)
                curr_para = ""

            if para.text.strip() == "": 
                if not in_para:
                    continue 
                else: # we are in a paragraph and we hit an empty line --> we are done with the paragraph
                    in_para = False
                    paras.append(curr_para)
                    curr_para = ""

            else: # not empty line --> we are in a paragraph
                if not in_para: 
                    in_para = True
                curr_para += para.text + "\n"
            
        return paras

    def get_text_from_docx(self):
        f = open(self.sop_path, "rb")
        doc = docx.Document(f)
        return "\n".join([para.text for para in doc.paragraphs])
    
    def extract_prompt(self, api_prompt: Regulatory_API_Prompt, context: Optional[str] = None):
        if context is None: 
            context = self.get_text_from_docx()

        api_prompt.user_prompt += context

        api_key = dotenv.get_key(".env", "ANTHROPIC_API_KEY")

        messages = [
            {"role": "user", "content": api_prompt.sys_prompt},
        ]

        if api_prompt.assistant_prompt is not None:
            messages.append({"role": "assistant", "content": api_prompt.assistant_prompt})

        client = anthropic.Anthropic(api_key=api_key)
        response = client.messages.create(
            model=api_prompt.model,
            max_tokens=api_prompt.max_tokens,
            system=api_prompt.sys_prompt,
            messages=messages
        )

        return response.content[0].text
        
    
    







    
